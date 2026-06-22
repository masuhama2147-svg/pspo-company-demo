# -*- coding: utf-8 -*-
"""
三福HD（P・SPO）社長向け 手渡し提案資料（Word/.docx）ビルダー。
python-docx で日本語フォント・見出し・表・箇条書きを整えて出力する。
出力: docs/三福HD様_ご提案資料_2026-06-22.docx
方針: 一次情報整合・誇張しない（データは“ポテンシャル”として記述）。硬めの常体/敬体。
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x16, 0x24, 0x3A)
ORANGE = RGBColor(0xC0, 0x5A, 0x12)
GREY = RGBColor(0x55, 0x5F, 0x6B)
JP_FONT = "Yu Gothic"          # 社長PCに標準搭載。無ければWordが代替
JP_FONT_E = "Yu Gothic"

doc = Document()

# 既定フォント（和欧）
normal = doc.styles["Normal"]
normal.font.name = JP_FONT
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), JP_FONT_E)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(4)

def _set_jp(run):
    run.font.name = JP_FONT
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), JP_FONT_E)

def h1(text):
    p = doc.add_heading(level=1)
    run = p.add_run(text); _set_jp(run)
    run.font.color.rgb = NAVY; run.font.size = Pt(16); run.font.bold = True
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    return p

def h2(text):
    p = doc.add_heading(level=2)
    run = p.add_run(text); _set_jp(run)
    run.font.color.rgb = NAVY; run.font.size = Pt(13); run.font.bold = True
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    return p

def para(text, bold=False, color=None, size=10.5, after=4, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text); _set_jp(run)
    run.font.bold = bold; run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    if align is not None: p.alignment = align
    return p

def lead(text):
    """段落リード（太字・濃色の小見出し文）"""
    return para(text, bold=True, color=ORANGE, size=11, after=3)

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text); _set_jp(run); run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 + 0.25*level)
    p.paragraph_format.space_after = Pt(2)
    return p

def num(text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text); _set_jp(run); run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), hexcolor); tcPr.append(sh)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(htext); _set_jp(r)
        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shade(hdr[i], "16243A")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(val); _set_jp(r); r.font.size = Pt(9.5)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(1)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def rule():
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr(); pb = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),"C9D2DD")
    pb.append(bot); pPr.append(pb)

# ============================ 表紙 ============================
for _ in range(3): doc.add_paragraph()
t = para("P・SPO 改革提案", size=26, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("〜「ジム運営会社」から「松山発・データ×AIプラットフォームカンパニー」へ〜",
     size=13, bold=True, color=ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
para("UXデザイン改革 ／ 第一者行動データのAI活用 ／ 新規BtoB・データマーケットプレイス事業",
     size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
for _ in range(6): doc.add_paragraph()
para("2026年6月22日", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, after=2)
para("ご提案先：三福ホールディングス株式会社 御中（株式会社P・SPOカンパニー）",
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("提案：愛媛大学・松山大学 学生BPOチーム（小山 拓斗 ほか）＋ AI実装チーム",
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
para("※ 本書の数値は学生BPO一次分析・公開情報に基づく概算であり、最新の公式数値とは差異があり得ます。",
     size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
doc.add_page_break()

# ======================= 0. エグゼクティブサマリー =======================
h1("0. エグゼクティブサマリー（要旨）")
para("P・SPOは、ジム・自習室・カフェ・サウナ・カラオケ・ピラティスなど30種類以上を月額一本で使える"
     "“多業態モデル”を、松山を中心に80店舗超・会員約1.6万人へと急成長させてきました（2016年 約350名から）。"
     "顔認証による入退館で、来館の事実データもすでに取得されています。", after=4)
lead("しかし、その「データ」が、ほとんど価値に変換されていません。")
bullet("混雑状況は“ライブカメラ画像”を会員が目視で判断する状態。来館データは入退館の認証にしか使われていない。")
bullet("アプリは外部ブラウザへ遷移し、パーソナライズも、データの会員還元（ランク・記録・特典）も実装されていない。")
bullet("＝「金脈（第一者の行動データ）の上に座っているのに、掘れていない」状態です。")
para("")
lead("本提案の主張は3点です。")
num("【守り】顧客体験（UX）で勝つ。混雑可視化・近い順・1アプリ完結・パーソナライズで、退会の多い学生・18歳・女性に効かせ、習慣化で継続を伸ばす。")
num("【攻め】取得済みデータをAIで価値化する。需要予測・離反予測・レコメンド・占有予測を自己学習で回し、データ・ネットワーク効果という“堀”を築く。")
num("【新規事業】データとAIモデルを外部へ販売するプラットフォーム／マーケットプレイスを立ち上げ、"
    "ジム会社から「データ×AIのプラットフォームカンパニー」へ転換する。")
para("")
para("私たちを入れていただければ、この全体像を1年でスケール基盤まで実装します（後述ロードマップ）。"
     "まずはお手元で、改革後の体験デモを実際に触ってご確認ください。", bold=True)
para("公開デモ： https://masuhama2147-svg.github.io/pspo-company-demo/", bold=True, color=ORANGE)
doc.add_page_break()

# ======================= 1. 現状のUXデザインの問題点 =======================
h1("1. 現状のUXデザインの問題点")
para("現行のP・SPOアプリ／Webを実地確認した上での、設計上の主要課題です。"
     "いずれも「あるはずの価値」を取りこぼしている、という性質の問題です。", after=4)

h2("1-1. “映像”を見せて“答え”を出していない")
bullet("店舗ページの混雑表現はライブカメラ画像（10分間隔）。会員は画像を見て“自分で”空き具合を解釈している。")
bullet("本来は『占有率ゲージ＋ゾーン別レベル＋あと何台空き』として、来店判断に直結する“答え”を返すべき。")

h2("1-2. アプリが“1つで完結”していない")
bullet("外部ブラウザへ遷移する設計で、体験が分断。会員証・予約・混雑・特典が一気通貫でない。")
bullet("『ホーム画面に追加＝ネイティブ風起動（PWA）』で、ワンタップ・ログイン維持・外部遷移ゼロにできる。")

h2("1-3. パーソナライズとデータ還元が皆無")
bullet("『あなたに合う空き時間・近い店舗・通い方』の提案が無い。来館データが会員に還元されていない。")
bullet("ランク・連続記録・利用連動割引といった“習慣化”の仕組み（Fogg/Hooked/自己決定理論）が未実装。")

h2("1-4. 退会構造への打ち手が無い")
bullet("退会“率”自体は居場所化で業界比むしろ良好。問題は学生コホートの過剰退会（18歳・女性が突出）。")
bullet("来館間隔の急減を検知し、先回りで効かせる仕組みが無い。1人の退会防止LTVは新規獲得CACを上回る。")

para("")
para("→ これらは「UXを直す」だけの話に見えて、実は“データの蛇口（アプリ）”を作る話です。"
     "アプリが使われるほどデータが増え、AIが賢くなる――次章以降の事業に直結します。", bold=True)
doc.add_page_break()

# ======================= 2. データが活用されていない現状 =======================
h1("2. 「データが取れているのに、使われていない」現状")
para("P・SPOは“多業態×顔認証×多店舗”という、国内でも稀有なデータ取得構造を既に持っています。"
     "それが活用されていないことこそ、最大の機会損失です。", after=4)

h2("2-1. すでに取得済みのデータ")
bullet("顔認証による入退館タイムスタンプ（来館頻度・時間帯・曜日）。")
bullet("会員属性（年齢・性別・プラン）。")
bullet("利用業態・店舗（ジム／自習室／カフェ／サウナ／カラオケ…のどこに来たか）。")

h2("2-2. しかし、現状の使い道は")
bullet("入退館の“認証”にしか使われていない（＝ゲートを開けるだけ）。")
bullet("経営の意思決定（出店・配置・販促・退会予防）にも、会員体験にも回っていない。")

h2("2-3. まだ取れていない（が、取れる）データ")
bullet("ゾーン別・設備別の占有（既存の顔認証／軽量IoTの拡張で取得可能）。滞在時間内訳。")
bullet("業態間の回遊（ジム→カフェ→自習室）、予約・キャンセル、アプリ内行動、オプション決済。")
para("")
para("→ つまり P・SPO は、松山という商圏で『生活者の暮らしのデジタルツイン』を構築できる位置にいます。"
    "これは、次章で述べるとおり“桁違いの資産”です。", bold=True)
doc.add_page_break()

# ======================= 3. AI生データの収集ポテンシャル =======================
h1("3. AI生データの収集ポテンシャル（どれだけ集められるか）")
para("AIの価値は、結局「どれだけ良いデータを、どれだけ持っているか」で決まります。"
     "P・SPOが集め得るデータは、量・質・希少性のすべてで突出しています。", after=4)

h2("3-1. 量：放っておいても毎日積み上がる")
bullet("規模：80店舗超 × 会員 約1.6万人 × 365日。")
bullet("概算：仮に会員の3割が週3回来館すると、入退館イベントは約2,000件/日 ＝ 年間およそ70万イベント。"
       "ゾーン・業態・滞在を加えれば、行動ログは年間数百万〜に膨らむ。")
bullet("これが“追加コストほぼゼロ”で、来館のたびに自動で蓄積される。")

h2("3-2. 質：AIが最も学習しやすい“行動系列”データ")
bullet("第一者（first-party）：本人同意ベースで自社が直接取得。第三者データのような推定が不要。")
bullet("本人紐付き・連続：同一IDの来館が時系列でつながる＝離反・習慣・LTVを“系列”としてモデル化できる。")
bullet("多業態クロス：1人の運動・学習・飲食・美容・娯楽を1IDで縫う。これは他に類を見ない。")

h2("3-3. 希少性：ポストCookie時代に“原本”を持つ強さ")
bullet("広告・データ規制で第三者データは縮小。アグリゲータが“推定で買う”ものを、P・SPOは“原本”で保有。")
bullet("国勢調査やアンケート（年単位・低頻度・想起バイアス）では取れない、連続・行動ベースの粒度。")
para("")
para("→ データ量×質×希少性が揃う事業者は、地方では極めて稀。"
    "これは「AIモデルの燃料」を独占的に握っている、ということです。", bold=True)
doc.add_page_break()

# ======================= 4. 海外の類似ビジネスモデル =======================
h1("4. 海外は「データ→AI→新収益」で勝っている")
para("フィットネス／多業態領域の先行各社は、行動データをAIに変え、継続率・稼働・新規B2B収益の源泉にしています。"
     "P・SPOは、これらの“いいとこ取り”ができる稀有なポジションにあります。", after=6)
table(
    ["企業・モデル", "データ → AIの活用", "P・SPOへの学び"],
    [
        ["Peloton（米）", "運動ログ→継続予測・レコメンド・コンテンツ最適化。ハード×サブスク×データの三位一体。",
         "継続率（churn）をデータで“商品の中心”に据える。"],
        ["ClassPass（米）", "予約・稼働データ→動的価格・需要予測で空き枠を最適化。",
         "稼働データ＝収益最適化エンジンになる。"],
        ["Planet Fitness（米）", "来館・混雑データ→出店立地・人員配置・解約予防。大規模・低価格でも“データ運営”。",
         "低価格・多店舗でもデータで運営できる。"],
        ["Whoop / Strava", "生体・行動データ→自己学習レコメンド。匿名集計を研究・自治体・保険へ二次提供。",
         "行動データの二次利用（B2B）が新収益になる。"],
        ["Gympass / Wellhub", "企業の健康経営×ジム利用データ→法人向けSaaS（BtoB）。",
         "健康経営は巨大なBtoB市場の入口。"],
        ["乐刻(Lefit)ほか（中）", "スマートジム・無人運営・稼働データ→AIで店舗運営を自動化し、FCを量産。",
         "“データ×無人×FC”で高速スケール（P・SPOと酷似）。"],
    ],
    widths=[1.5, 3.2, 2.1],
)
para("共通点は明確です。勝者は『データを集める仕組み（アプリ）→AIで価値化→新しい収益（B2B/FC）』を回しています。"
     "P・SPOは“多業態×無人×不動産シナジー”を併せ持つ点で、むしろ各社より有利です。", bold=True)
doc.add_page_break()

# ======================= 5. AI研究者の視点 =======================
h1("5. AI研究者の視点：自己学習するAIモデル群")
para("取得した行動データの上に、次のAIモデルを段階的に構築します。"
     "重要なのは、人手でラベル付けせずとも“自己学習（self-supervised）”で賢くなる設計です。", after=4)

h2("5-1. 構築するAIモデル")
bullet("離反予測（churn）：来館間隔の急減を早期検知。18歳・女性・学生に先回り特典（LTV>CAC）。")
bullet("需要・占有予測：店舗×時間×業態の混雑を予測。会員へ『あと何台空き』、経営へ人員配置最適化。")
bullet("レコメンド：『あなたに合う空き時間・近い店舗・通い方』を非PIIで提示し、習慣化を後押し。")
bullet("行動系列モデル（自己教師あり）：来館シーケンスから次回来館・LTVを予測。ラベル不要で自律的に改善。")
bullet("LLMエージェント：会員サポート・予約・FAQの自動化。社内は“自然言語でデータに質問”するBIエージェント。")

h2("5-2. なぜ“堀（moat）”になるのか")
bullet("データ・ネットワーク効果：アプリが使われる→データ増→AIが賢く→体験が良くなる→さらに使われる、の好循環。")
bullet("先行者優位：第一者・多業態・連続データの“原本”を松山で持つのはP・SPOだけ。後発は追随困難。")
bullet("データクリーンルーム／連合学習：非PIIのまま外部と安全に協業し、価値を広げられる。")
para("")
para("→ AIは“導入して終わり”ではなく、使うほど価値が増える資産になります。これが模倣困難性の正体です。", bold=True)
doc.add_page_break()

# ======================= 6. 新ビジネスモデル：プラットフォーム化 =======================
h1("6. 新ビジネスモデル：BtoB・プラットフォームカンパニーへ")
para("P・SPOの本質は、もはや「ジム運営会社」ではなく「松山の生活者データを握るプラットフォーム」です。"
     "そこから生まれるBtoB事業を、データ資産を核に展開します。", after=4)

h2("6-1. 展開する事業（B2C収益に“載せる”）")
num("人流・行動データのB2B提供：小売・商店街・観光・デベロッパー・自治体へ。匿名加工・同意ベース・店舗単位集計。")
num("健康経営パッケージ：法人会員×健康・利用データのレポート。健康経営という巨大B2B市場の入口。")
num("出店立地AI：空き店舗×需要をスコアリング。三福綜合不動産の“シャッターゼロ”を高速化し、外部にも提供。")
num("多業態運営パッケージ（FC）：多業態×無人×データ運営を“型”にして他地域・他社へ。のれん分けカンパニー制と整合。")
num("AIモデル／データのマーケットプレイス（次章）。")

h2("6-2. 何が変わるか")
bullet("収益が「店舗の月会費（B2C・労働集約）」一本足から、「データ・AI（B2B・高利益・スケール）」へ多角化。")
bullet("地方の一事業者から、“松山発のデータプラットフォーム”という全国・全業種に売れる立ち位置へ。")
doc.add_page_break()

# ======================= 7. マーケットプレイス構想 =======================
h1("7. マーケットプレイス構想：P・SPO Data & AI Marketplace")
para("スケールの主役は「マーケットプレイス」です。P・SPOが蓄積した行動データと、その上に構築したAIモデルを、"
     "外部が安全に利用できるプラットフォームを立ち上げます。", after=4)

h2("7-1. 出品物（何を売るか）")
bullet("匿名加工データセット：人流・属性別来館・商圏トレンド（店舗単位・同意ベース）。")
bullet("学習済みAIモデル：需要予測・離反予測・出店立地スコアリング・混雑予測を“すぐ使える形”で。")
bullet("API／ダッシュボード：自社データに重ねて使える分析基盤。")
bullet("多業態運営パッケージ：FC・他事業者向けの“型”（運営×データ×AIのセット）。")

h2("7-2. 買い手（誰に売るか）")
bullet("小売・飲食チェーン（出店・販促・需要予測）／デベロッパー（テナント・再開発）。")
bullet("自治体・県（EBPM・回遊・観光・シャッターゼロの効果測定）。")
bullet("保険・ヘルスケア（健康行動データの活用）／他のフィットネス・多業態事業者（FC）。")

h2("7-3. なぜP・SPOが勝てるか")
bullet("“原本”を持つ唯一性＋データ・ネットワーク効果でロックイン。先行するほど差が開く。")
bullet("収益モデル：データ販売／モデルのサブスク・従量／API／FCライセンスの複線。")
para("")
para("→ これが、ジム会社を「データ×AIのプラットフォームカンパニー」へ変える“最後のピース”です。", bold=True)
doc.add_page_break()

# ======================= 8. 1年スケーリング・ロードマップ =======================
h1("8. もし入れていただけたら：1年スケーリング・ロードマップ")
para("私たちが参画した場合の、12ヶ月での実装計画です。各フェーズで“測れるKPI”を出しながら進めます。", after=6)
table(
    ["フェーズ（期間）", "やること", "成果・KPI"],
    [
        ["Phase 0：本提案（完了）", "UX改修デモ＋データ価値の可視化（本日の公開サイト）。",
         "改革後の体験を実機で確認。"],
        ["Phase 1：0–3ヶ月", "P・SPOアプリMVP（混雑可視化・近い順・パーソナライズ・1アプリ完結）＋データ基盤（イベント収集・非PII射影・同意管理）。学生離反対策を先行投入。",
         "アプリ稼働。来館データのパイプライン化。学生継続率の改善着手。"],
        ["Phase 2：3–6ヶ月", "社内AIモデル（離反・需要/占有予測・レコメンド）＋BIエージェント。",
         "継続率・稼働の改善を数値で実証。意思決定のデータ化。"],
        ["Phase 3：6–9ヶ月", "B2B商品化（人流データレポート・健康経営・出店立地AI）。最初の外部PoC。",
         "外部顧客（自治体/小売/デベロッパー）で初収益。"],
        ["Phase 4：9–12ヶ月", "データ/AIマーケットプレイスβ＋FCパッケージ。",
         "データ・AIモデルの外部提供開始。プラットフォーム事業の初収益。"],
    ],
    widths=[1.7, 3.3, 1.8],
)
para("12ヶ月後：「ジム運営」＋「データ×AIプラットフォーム」の二本柱と、全国・全業種へ売れるスケール基盤が完成します。", bold=True)
doc.add_page_break()

# ======================= 9. 本日お渡しするデモ各ページの説明 =======================
h1("9. 本日のデモ各ページの説明（何を示すか）")
para("公開URL： https://masuhama2147-svg.github.io/pspo-company-demo/ "
     "（ビルド済みの静的サイト。スマホ・PC両対応。検索には出さない設定）", color=ORANGE, after=6)
pages = [
    ("トップ（index）", "全体の入口。役割別の導線＋公式リール（実写）＋実績。"
                      "“ジム会社”から“データプラットフォーム”への世界観を提示。"),
    ("日常アプリ（app）", "会員の毎日の画面。『今、空いてる？』をひと目で。1アプリ完結のUX。"),
    ("店舗・空き（area）", "80店舗超を地図＋『近い順』で。データの可視化と回遊設計。"),
    ("混雑ライブ（live）", "占有ゲージ＝“カメラ画像”を“答え（あと何台空き）”に変える提案。"),
    ("マイページ（me）", "来館データを会員に還元（ランク・連続記録）＝習慣化の仕掛け。"),
    ("料金まるわかり（price）", "複雑な料金を、Apple製品ページ風のビルダーで明快に。"),
    ("はじめ方（guide）", "入会の摩擦を下げるオンボーディング。"),
    ("経営×成長（growth）", "データ→意思決定。離反・稼働・出店の経営ダッシュボード思想。"),
    ("データ活用（data）", "取れるデータと、その二次利用（B2B/EBPM）の地図。新規事業の核。"),
    ("まちみる／B2B（machimiru）", "人流データを地域・企業へ。プラットフォーム事業の入口。"),
    ("提案サマリ（brief）", "本提案の要約（経営層向けの1枚絵）。"),
    ("自己批評（critique）", "現状の弱点を誠実に直視。誇張せず、信頼を担保するための章。"),
]
for name, desc in pages:
    p = doc.add_paragraph()
    r = p.add_run("● " + name + "： "); _set_jp(r); r.font.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(10.5)
    r2 = p.add_run(desc); _set_jp(r2); r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
doc.add_page_break()

# ======================= 10. クロージング =======================
h1("10. ご提案のまとめ（クロージング）")
lead("P・SPOは、すでにAI時代の“燃料（第一者・多業態・連続の行動データ）”を握っています。")
para("足りないのは、それを①体験（UX）で増やし、②AIで価値化し、③外部へ売る（プラットフォーム化）――この3本の実装だけです。",
     after=6)
para("私たちは、UX設計・データ基盤・AIモデル・BtoB事業化までを一気通貫で実装できるチームです。"
     "まずは Phase 1（アプリMVP＋データ基盤）から着手し、1年でスケール基盤までお持ちします。", after=6)
lead("本日のお願い：Phase 1 の実装機会をいただきたく存じます。")
para("まずは、お手元のスマホ・PCで公開デモに触れて、改革後の“体験の差”をご確認ください。", after=8)
para("公開デモ： https://masuhama2147-svg.github.io/pspo-company-demo/", bold=True, color=ORANGE)
para("GitHub（全ソース・設計資料）： https://github.com/masuhama2147-svg/pspo-company-demo", color=GREY, size=9.5)
rule()
para("提案：愛媛大学・松山大学 学生BPOチーム（小山 拓斗 ほか）＋ AI実装チーム／ 2026年6月22日",
     size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.RIGHT)

import os
os.makedirs("docs", exist_ok=True)
out = "docs/三福HD様_ご提案資料_2026-06-22.docx"
doc.save(out)
print("✓ 生成:", out)
